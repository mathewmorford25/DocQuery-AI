import os
import tempfile
from dotenv import load_dotenv

import streamlit as st
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from docx import Document as DocxDocument


load_dotenv()

def get_openai_api_key():
    """
    Get OpenAI API key from Streamlit secrets (for deployment)
    or from .env (for local development).
    """
    try:
        return st.secrets["OPENAI_API_KEY"]
    except Exception:
        return os.getenv("OPENAI_API_KEY")

def deduplicate_docs(docs: list[Document]) -> list[Document]:
    """
    Remove duplicate retrieved docs based on source file, page, and content.
    """
    unique_docs = []
    seen = set()

    for doc in docs:
        key = (
            doc.metadata.get("source_file", "Unknown file"),
            doc.metadata.get("page", None),
            doc.page_content.strip()
        )

        if key not in seen:
            seen.add(key)
            unique_docs.append(doc)

    return unique_docs

def load_documents(uploaded_files) -> list[Document]:
    """
    Load uploaded PDF, DOCX, and TXT files into LangChain Document objects.
    """
    documents = []

    for uploaded_file in uploaded_files:
        file_name = uploaded_file.name
        file_extension = file_name.lower().split(".")[-1]

        if file_extension == "pdf":
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
                temp_file.write(uploaded_file.read())
                temp_path = temp_file.name

            loader = PyPDFLoader(temp_path)
            file_docs = loader.load()

            for doc in file_docs:
                doc.metadata["source_file"] = file_name

            documents.extend(file_docs)

            try:
                os.remove(temp_path)
            except OSError:
                pass

        elif file_extension == "docx":
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
                temp_file.write(uploaded_file.read())
                temp_path = temp_file.name

            docx_file = DocxDocument(temp_path)
            full_text = "\n".join(
                [paragraph.text for paragraph in docx_file.paragraphs if paragraph.text.strip()]
            )

            documents.append(
                Document(
                    page_content=full_text,
                    metadata={
                        "source_file": file_name,
                        "page": 0
                    }
                )
            )

            try:
                os.remove(temp_path)
            except OSError:
                pass

        elif file_extension == "txt":
            text_content = uploaded_file.read().decode("utf-8", errors="ignore")

            documents.append(
                Document(
                    page_content=text_content,
                    metadata={
                        "source_file": file_name,
                        "page": 0
                    }
                )
            )

    return documents


def chunk_documents(
    documents: list[Document],
    chunk_size: int,
    chunk_overlap: int
) -> list[Document]:
    """
    Split long documents into smaller overlapping chunks.
    """
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap
    )
    return text_splitter.split_documents(documents)


def build_vectorstore(chunks: list[Document], api_key: str) -> FAISS:
    """
    Convert chunks to embeddings and store them in FAISS.
    """
    embeddings = OpenAIEmbeddings(api_key=api_key)
    vectorstore = FAISS.from_documents(chunks, embeddings)
    return vectorstore


def format_sources(source_docs: list[Document]) -> list[str]:
    """
    Group sources by file and combine page numbers into a cleaner citation list.
    """
    grouped = {}

    for doc in source_docs:
        source_file = doc.metadata.get("source_file", "Unknown file")
        page = doc.metadata.get("page", None)

        if source_file not in grouped:
            grouped[source_file] = set()

        if page is not None:
            grouped[source_file].add(page + 1)

    formatted = []

    for source_file, pages in grouped.items():
        if pages:
            sorted_pages = sorted(pages)
            page_text = ", ".join(str(p) for p in sorted_pages)
            formatted.append(f"{source_file} — pages {page_text}")
        else:
            formatted.append(source_file)

    return formatted


def answer_question(
    question: str,
    vectorstore: FAISS,
    retrieval_k: int,
    chat_history: list[dict],
    api_key: str
):
    """
    Retrieve relevant chunks and ask the model to answer using only that context.
    """
    retriever = vectorstore.as_retriever(search_kwargs={"k": retrieval_k})
    relevant_docs = retriever.invoke(question)

    recent_history = chat_history[-3:] if chat_history else []

    conversation_text = "\n\n".join(
        [
            f"Previous Question: {item['question']}\nPrevious Answer: {item['answer']}"
            for item in recent_history
        ]
    )

    context_text = "\n\n".join(
        [
            f"Source: {doc.metadata.get('source_file', 'Unknown')} | "
            f"Page: {doc.metadata.get('page', 'N/A')}\n"
            f"{doc.page_content}"
            for doc in relevant_docs
        ]
    )

    prompt = f"""
    You are a helpful corporate document assistant.

    Use the recent conversation history for context when the user's question is a follow-up.
    Answer the user's question using ONLY the retrieved document context below.
    If the answer is not in the context, say:
    "I could not find that in the uploaded documents."

    Do not include a citations section.
    Do not mention sources in the answer.
    Just provide a clear, concise answer.

    Recent Conversation History:
    {conversation_text}

    Retrieved Document Context:
    {context_text}

    Current Question:
    {question}
    """

    llm = ChatOpenAI(
        model="gpt-4.1-mini",
        temperature=0,
        api_key=api_key
    )
    response = llm.invoke(prompt)

    return response.content, relevant_docs

def preview_text(text: str, max_chars: int = 400) -> str:
    """
    Return a shortened preview of long text for cleaner UI display.
    """
    cleaned = " ".join(text.split())
    if len(cleaned) <= max_chars:
        return cleaned
    return cleaned[:max_chars].rstrip() + "..."

def build_chat_export(chat_history: list[dict]) -> str:
    """
    Build a text export of the conversation history.
    """
    lines = []
    lines.append("Corporate RAG Assistant - Conversation Export")
    lines.append("=" * 50)
    lines.append("")

    for idx, item in enumerate(chat_history, start=1):
        lines.append(f"Question {idx}: {item['question']}")
        lines.append("")
        lines.append("Answer:")
        lines.append(item["answer"])
        lines.append("")
        lines.append("Sources Used:")
        for source in item["sources"]:
            lines.append(f"- {source}")
        lines.append("")
        lines.append("-" * 50)
        lines.append("")

    return "\n".join(lines)

def get_suggested_questions() -> list[str]:
    """
    Return helpful starter questions for the user.
    """
    return [
        "What is this document about?",
        "Who is eligible?",
        "What are the requirements?",
        "What documents are required?",
        "Does it mention background checks?",
        "Summarize this in plain English."
    ]

def main():
    st.set_page_config(page_title="Corporate RAG Assistant", layout="wide")
    st.title("DocQuery AI")
    st.caption("Upload multiple PDF documents, process them into searchable chunks, and ask grounded questions with source-backed answers.")

    api_key = get_openai_api_key()

    if not api_key:
        st.error("OPENAI_API_KEY is missing.")
        st.stop()

    if "vectorstore" not in st.session_state:
        st.session_state.vectorstore = None

    if "documents_processed" not in st.session_state:
        st.session_state.documents_processed = False

    if "last_answer" not in st.session_state:
        st.session_state.last_answer = ""

    if "last_source_docs" not in st.session_state:
        st.session_state.last_source_docs = []

    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []

    if "selected_question" not in st.session_state:
        st.session_state.selected_question = ""

    # =========================
    # Sidebar Settings
    # =========================
    st.sidebar.title("Settings")
    st.sidebar.caption("Adjust how the assistant processes and retrieves document content.")

    chunk_size = st.sidebar.slider(
        "Chunk Size",
        min_value=300,
        max_value=2000,
        value=1000,
        step=100,
        help="Controls how much text goes into each chunk."
    )

    chunk_overlap = st.sidebar.slider(
        "Chunk Overlap",
        min_value=0,
        max_value=500,
        value=150,
        step=25,
        help="Controls how much overlap exists between chunks."
    )

    retrieval_k = st.sidebar.slider(
        "Retrieved Chunks",
        min_value=1,
        max_value=10,
        value=4,
        step=1,
        help="Controls how many chunks are retrieved to answer a question."
    )

    st.sidebar.markdown("---")
    st.sidebar.markdown("**Current Configuration**")
    st.sidebar.write(f"Chunk Size: {chunk_size}")
    st.sidebar.write(f"Chunk Overlap: {chunk_overlap}")
    st.sidebar.write(f"Retrieved Chunks: {retrieval_k}")

    # =========================
    # Document Setup Section
    # =========================
    with st.container():
        st.subheader("Document Setup")
        st.caption("Upload PDF, DOCX, or TXT files, process them, and prepare the assistant for questions.")

        uploaded_files = st.file_uploader(
            "Upload document files",
            type=["pdf", "docx", "txt"],
            accept_multiple_files=True
        )

        # Reset app state if all files are removed
        if not uploaded_files:
            st.session_state.documents_processed = False
            st.session_state.vectorstore = None
            st.session_state.last_answer = ""
            st.session_state.last_source_docs = []
            st.session_state.chat_history = []

        # Clear old results when new files are uploaded but not processed yet
        if uploaded_files and not st.session_state.documents_processed:
            st.session_state.last_answer = ""
            st.session_state.last_source_docs = []
            st.session_state.chat_history = []

        if st.button("Process Documents"):
            if not uploaded_files:
                st.warning("Please upload at least one PDF.")
            else:
                with st.spinner("Loading documents..."):
                    docs = load_documents(uploaded_files)

                with st.spinner("Chunking documents..."):
                    chunks = chunk_documents(docs, chunk_size, chunk_overlap)

                with st.spinner("Building vector index..."):
                    st.session_state.vectorstore = build_vectorstore(chunks, api_key)

                st.session_state.documents_processed = True
                st.session_state.last_answer = ""
                st.session_state.last_source_docs = []
                st.session_state.chat_history = []

                st.success(
                    f"Documents processed successfully. Indexed {len(chunks)} chunks from {len(uploaded_files)} file(s)."
                )

        if st.session_state.documents_processed:
            st.success("Documents are processed and ready for questions.")

    # =========================
    # Ask the Assistant Section
    # =========================
    with st.container():
        st.subheader("Ask the Assistant")
        st.caption("Ask questions about the uploaded documents.")

        question = st.text_input(
            "Ask a question about the uploaded documents",
            value=st.session_state.selected_question,
            disabled=not st.session_state.documents_processed,
            placeholder="Example: What is this document about?"
        )

        if st.button("Ask", disabled=not st.session_state.documents_processed):
            if not st.session_state.vectorstore:
                st.warning("Please upload and process documents first.")
            elif not question.strip():
                st.warning("Please enter a question.")
            else:
                with st.spinner("Searching documents and generating answer..."):
                    answer, source_docs = answer_question(
                        question,
                        st.session_state.vectorstore,
                        retrieval_k,
                        st.session_state.chat_history,
                        api_key
                    )

                source_docs = deduplicate_docs(source_docs)

                st.session_state.last_answer = answer
                st.session_state.last_source_docs = source_docs

                st.session_state.chat_history.append(
                    {
                        "question": question,
                        "answer": answer,
                        "sources": format_sources(source_docs),
                        "context_docs": source_docs
                    }
                )

                st.session_state.selected_question = ""

    # =========================
    # Results Section
    # =========================
    with st.container():
        if st.session_state.chat_history:
            st.markdown("---")
            st.subheader("Results")

            source_count = len({source for item in st.session_state.chat_history for source in item["sources"]})
            latest_context_count = len(st.session_state.chat_history[-1]["context_docs"])

            col1, col2 = st.columns(2)
            with col1:
                st.metric("Source Files", source_count)
            with col2:
                st.metric("Retrieved Chunks", latest_context_count)

            export_text = build_chat_export(st.session_state.chat_history)

            col3, col4 = st.columns(2)
            with col3:
                st.download_button(
                    label="Download Conversation",
                    data=export_text,
                    file_name="conversation_history.txt",
                    mime="text/plain"
                )
            with col4:
                if st.button("Clear Conversation"):
                    st.session_state.last_answer = ""
                    st.session_state.last_source_docs = []
                    st.session_state.chat_history = []
                    st.rerun()

            st.markdown("### Conversation History")

            for idx, item in enumerate(reversed(st.session_state.chat_history), start=1):
                question_number = len(st.session_state.chat_history) - idx + 1

                with st.expander(f"Question {question_number}: {item['question']}", expanded=(idx == 1)):
                    st.markdown("**Answer**")
                    st.info(item["answer"])

                    st.markdown("**Sources Used**")
                    for source in item["sources"]:
                        st.write(f"- {source}")

                    st.markdown("**Retrieved Context**")
                    for i, doc in enumerate(item["context_docs"], start=1):
                        source_file = doc.metadata.get("source_file", "Unknown file")
                        page = doc.metadata.get("page", None)

                        if page is not None:
                            label = f"{source_file} — page {page + 1}"
                        else:
                            label = source_file

                        with st.expander(f"Context {i}: {label}"):
                            st.markdown("**Preview**")
                            st.write(preview_text(doc.page_content, max_chars=400))

                            st.markdown("**Full Chunk Text**")
                            st.code(doc.page_content, language=None)
if __name__ == "__main__":
    main()

st.sidebar.markdown("---")
st.sidebar.markdown("### About")
st.sidebar.write(
    "DocQuery AI is a multi-document RAG app that lets users upload PDF, DOCX, and TXT files, ask grounded questions, and receive source-backed answers with retrieved context."
)
st.sidebar.write("Built by Mat Morford")